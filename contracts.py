"""
contracts.py
============
Generator dokumen kontrak kerja (PKWT) otomatis berdasarkan data kandidat &
posisi. Menghasilkan file PDF (reportlab) dan DOCX (python-docx) yang siap
diunduh / ditandatangani digital.
"""

import os
from datetime import datetime, timedelta

import config
import db
from catalog import job_by_id


def _contract_context(doc: dict) -> dict:
    job = job_by_id(doc["jobId"]) or {}
    today = datetime.now()
    end = today + timedelta(days=30 * 6)  # PKWT 6 bulan (dapat dikonfigurasi)
    is_maritime = doc["track"] == "maritime"
    if is_maritime:
        salary = f"USD {int(job.get('salaryMaxUsd', 0)):,}".replace(",", ".")
        dept = job.get("vesselHint", "Maritime Operations")
        jalur = "Maritime Crew"
    else:
        salary = f"Rp{int(job.get('salaryMaxIdr', 0)):,}".replace(",", ".")
        dept = job.get("dept", "Back Office")
        jalur = "Back Office Administration"
    no_kontrak = f"LMI/PKWT/{today.year}/{db.now_iso()[-6:].replace(':', '').replace('-', '')[:4]}"
    return {
        "no_kontrak": no_kontrak, "today": today, "end": end, "is_maritime": is_maritime,
        "salary": salary, "dept": dept, "jalur": jalur, "job": job,
        "name": doc.get("name", ""), "jobTitle": doc.get("jobTitle", ""),
        "cert_number": (doc.get("cert") or {}).get("number"),
    }


def _fmt(d: datetime) -> str:
    bulan = ["Januari", "Februari", "Maret", "April", "Mei", "Juni",
             "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    return f"{d.day} {bulan[d.month - 1]} {d.year}"


# ---------------------------------------------------------------------------
# PDF (reportlab)
# ---------------------------------------------------------------------------

def generate_pdf(doc: dict) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    from reportlab.lib import colors

    ctx = _contract_context(doc)
    path = os.path.join(config.CONTRACT_DIR, f"PKWT_{doc['id']}.pdf")

    styles = getSampleStyleSheet()
    h_company = ParagraphStyle("company", parent=styles["Title"], fontSize=16, spaceAfter=2,
                               textColor=colors.HexColor("#0F2A33"))
    h_sub = ParagraphStyle("sub", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#5b6b73"),
                           spaceAfter=14)
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, leading=16, alignment=TA_JUSTIFY,
                          textColor=colors.HexColor("#26343b"))
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#5b6b73"))
    sign = ParagraphStyle("sign", parent=styles["Normal"], fontSize=9, alignment=TA_CENTER)

    pdf = SimpleDocTemplate(path, pagesize=A4, topMargin=22 * mm, bottomMargin=22 * mm,
                            leftMargin=20 * mm, rightMargin=20 * mm,
                            title=f"PKWT {ctx['name']}")
    story = []
    story.append(Paragraph(config.COMPANY_LEGAL_NAME.upper(), h_company))
    story.append(Paragraph("Perjanjian Kerja Waktu Tertentu (PKWT)", h_sub))
    story.append(Paragraph(f"No: {ctx['no_kontrak']} &nbsp;&nbsp;&middot;&nbsp;&nbsp; {_fmt(ctx['today'])}", small))
    story.append(Spacer(1, 12))

    story.append(Paragraph(
        f"Perjanjian Kerja Waktu Tertentu ini dibuat antara <b>{config.COMPANY_LEGAL_NAME}</b> "
        f"(selanjutnya disebut &ldquo;Perusahaan&rdquo;) dan pihak Pekerja yang identitasnya tercantum "
        f"di bawah ini, dengan ketentuan sebagai berikut:", body))
    story.append(Spacer(1, 12))

    rows = [
        ["Nama Lengkap", ctx["name"]],
        ["Jabatan / Posisi", ctx["jobTitle"]],
        ["Departemen / Unit", ctx["dept"]],
        ["Jalur Rekrutmen", ctx["jalur"]],
        ["Jenis Kontrak", "PKWT — Waktu Tertentu (6 Bulan)"],
        ["Tanggal Mulai", _fmt(ctx["today"])],
        ["Tanggal Berakhir", _fmt(ctx["end"])],
        ["Gaji Pokok", f"{ctx['salary']} / Bulan"],
    ]
    if ctx["is_maritime"] and ctx["cert_number"]:
        rows.append(["No. Sertifikat Pelaut", ctx["cert_number"]])

    tbl = Table([[Paragraph(f"<b>{k}</b>", small), Paragraph(str(v), small)] for k, v in rows],
                colWidths=[55 * mm, None])
    tbl.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.HexColor("#e2e8f0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 16))

    story.append(Paragraph(
        "Pekerja bersedia mematuhi seluruh peraturan Perusahaan, menjaga kerahasiaan informasi, "
        "serta melaksanakan tugas sesuai standar kompetensi yang dipersyaratkan. Kontrak dapat "
        "diperpanjang berdasarkan evaluasi kinerja dan kebutuhan operasional Perusahaan.", body))
    story.append(Spacer(1, 30))

    sig_data = [[
        Paragraph(f"Pekerja<br/><br/><br/><br/><b>{ctx['name']}</b>", sign),
        Paragraph(f"HR Manager<br/><br/><br/><br/><b>{config.COMPANY_SHORT_NAME}</b>", sign),
        Paragraph(f"Direktur Utama<br/><br/><br/><br/><b>{config.COMPANY_SHORT_NAME}</b>", sign),
    ]]
    sig = Table(sig_data, colWidths=[None, None, None])
    sig.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(Paragraph(f"{config.COMPANY_CITY}, {_fmt(ctx['today'])}", small))
    story.append(Spacer(1, 8))
    story.append(sig)

    pdf.build(story)
    return path


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------

def generate_docx(doc: dict) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ctx = _contract_context(doc)
    path = os.path.join(config.CONTRACT_DIR, f"PKWT_{doc['id']}.docx")
    d = Document()

    title = d.add_paragraph()
    run = title.add_run(config.COMPANY_LEGAL_NAME.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0F, 0x2A, 0x33)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = d.add_paragraph("Perjanjian Kerja Waktu Tertentu (PKWT)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = d.add_paragraph(f"No: {ctx['no_kontrak']}  ·  {_fmt(ctx['today'])}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    d.add_paragraph(
        f"Perjanjian Kerja Waktu Tertentu ini dibuat antara {config.COMPANY_LEGAL_NAME} "
        f"(selanjutnya disebut \u201cPerusahaan\u201d) dan pihak Pekerja yang identitasnya tercantum "
        f"di bawah ini, dengan ketentuan sebagai berikut:")

    rows = [
        ("Nama Lengkap", ctx["name"]),
        ("Jabatan / Posisi", ctx["jobTitle"]),
        ("Departemen / Unit", ctx["dept"]),
        ("Jalur Rekrutmen", ctx["jalur"]),
        ("Jenis Kontrak", "PKWT — Waktu Tertentu (6 Bulan)"),
        ("Tanggal Mulai", _fmt(ctx["today"])),
        ("Tanggal Berakhir", _fmt(ctx["end"])),
        ("Gaji Pokok", f"{ctx['salary']} / Bulan"),
    ]
    if ctx["is_maritime"] and ctx["cert_number"]:
        rows.append(("No. Sertifikat Pelaut", ctx["cert_number"]))

    table = d.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)

    d.add_paragraph()
    d.add_paragraph(
        "Pekerja bersedia mematuhi seluruh peraturan Perusahaan, menjaga kerahasiaan informasi, "
        "serta melaksanakan tugas sesuai standar kompetensi yang dipersyaratkan. Kontrak dapat "
        "diperpanjang berdasarkan evaluasi kinerja dan kebutuhan operasional Perusahaan.")
    d.add_paragraph()
    d.add_paragraph(f"{config.COMPANY_CITY}, {_fmt(ctx['today'])}")
    d.add_paragraph()

    sig = d.add_table(rows=1, cols=3)
    labels = ["Pekerja", "HR Manager", "Direktur Utama"]
    names = [ctx["name"], config.COMPANY_SHORT_NAME, config.COMPANY_SHORT_NAME]
    for i in range(3):
        cell = sig.rows[0].cells[i]
        cell.text = labels[i] + "\n\n\n\n" + names[i]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d.save(path)
    return path


def generate_contract(maxy_id: str, fmt: str = "pdf") -> str | None:
    """Generate kontrak untuk kandidat. fmt = 'pdf' | 'docx'. Return path file."""
    doc = db.get_candidate(maxy_id)
    if not doc:
        return None
    path = generate_pdf(doc) if fmt == "pdf" else generate_docx(doc)
    doc["contractGenerated"] = True
    db.save_candidate(doc)
    return path
